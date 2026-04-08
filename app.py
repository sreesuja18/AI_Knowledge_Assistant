import streamlit as st
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document
from groq import Groq
from rank_bm25 import BM25Okapi
from deep_translator import GoogleTranslator
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from io import BytesIO
import os
import time
import pickle

st.set_page_config(page_title="AI Knowledge Assistant", page_icon="🤖", layout="centered")

st.markdown("""
<style>
.welcome-box {
    background: linear-gradient(135deg, #1F4E79, #2E75B6);
    color: white; padding: 1.2rem 1.5rem;
    border-radius: 12px; margin-bottom: 1.2rem;
}
.welcome-box h3 { margin: 0 0 4px 0; font-size: 18px; }
.welcome-box p  { margin: 0; font-size: 13px; opacity: 0.88; }
.company-tag {
    display: inline-block; background: #E3F2FD; color: #0D47A1;
    border-radius: 12px; padding: 2px 10px;
    font-size: 11px; font-weight: 600; margin-right: 6px;
}
.chunk-preview {
    border-left: 3px solid #2E75B6; padding: 6px 10px;
    font-size: 12px; color: #555; margin-top: 4px;
    border-radius: 0 6px 6px 0; font-style: italic;
}
.confidence-high { color: #1B5E20; font-weight: 600; font-size: 12px; }
.confidence-med  { color: #E65100; font-weight: 600; font-size: 12px; }
.confidence-low  { color: #B71C1C; font-weight: 600; font-size: 12px; }
.time-badge {
    display: inline-block; background: #F3E5F5; color: #6A1B9A;
    border-radius: 10px; padding: 2px 8px;
    font-size: 11px; font-weight: 500; margin-left: 8px;
}
.tech-badge {
    display: inline-block; background: #E8F5E9; color: #1B5E20;
    border-radius: 10px; padding: 2px 8px;
    font-size: 11px; font-weight: 500; margin-left: 4px;
}
</style>
""", unsafe_allow_html=True)

@st.cache_resource
def load_resources():
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    vector_db = FAISS.load_local("vector_db", embeddings, allow_dangerous_deserialization=True)
    api_key = None
    try:
        api_key = st.secrets["GROQ_API_KEY"]
    except Exception:
        api_key = "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"
    client = Groq(api_key=api_key)
    with open("bm25_chunks.pkl", "rb") as f:
        data = pickle.load(f)
    tokenized = [text.lower().split() for text in data["texts"]]
    bm25 = BM25Okapi(tokenized)
    return vector_db, client, bm25, data["texts"], data["metas"]

vector_db, client, bm25, chunk_texts, chunk_metas = load_resources()

def count_docs():
    try:
        return len([f for f in os.listdir("data") if f.endswith(".pdf")])
    except:
        return 0

def get_company(filename):
    name = os.path.basename(filename).replace("_", " ").replace(".pdf", "")
    companies = ["TCS", "Infosys", "Wipro", "Cognizant", "Google", "Tesla",
                 "Amazon", "Microsoft", "Meta", "Johnson"]
    for c in companies:
        if c.lower() in name.lower():
            return "Johnson & Johnson" if c == "Johnson" else c
    return name[:25]

def translate_text(text, target_lang):
    try:
        return GoogleTranslator(source="auto", target=target_lang).translate(text)
    except:
        return text

def get_llm_answer(question, context, history=[]):
    history_text = ""
    if history:
        history_text = "Previous conversation:\n"
        for h in history[-3:]:
            history_text += f"Q: {h['question']}\nA: {h['answer'][:200]}...\n\n"
    prompt = f"""You are a company HR policy assistant.
{history_text}
Answer using ONLY the context below. If the answer is not in the context,
say: "I could not find this information in the company documents."
When possible, present information in a structured format with bullet points or tables.

Context:
{context}

Current Question: {question}
Answer:"""
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content

def expand_query(question):
    prompt = f"""Rephrase this HR policy question in 2 different ways to improve search results.
Return only the 2 rephrased questions, one per line, nothing else.
Question: {question}"""
    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=100
        )
        extras = response.choices[0].message.content.strip().split("\n")
        extras = [e.strip() for e in extras if e.strip()]
        return [question] + extras[:2]
    except:
        return [question]

def hybrid_search(query, k=5):
    vector_results = vector_db.similarity_search_with_score(query, k=k)
    tokens = query.lower().split()
    bm25_scores = bm25.get_scores(tokens)
    top_bm25_idx = sorted(range(len(bm25_scores)),
                          key=lambda i: bm25_scores[i], reverse=True)[:k]
    combined = {}
    for doc, score in vector_results:
        combined[doc.page_content] = {"doc": doc, "score": 1 - min(score, 1)}
    for idx in top_bm25_idx:
        text = chunk_texts[idx]
        bm25_norm = bm25_scores[idx] / (max(bm25_scores) + 1e-9)
        if text in combined:
            combined[text]["score"] = (combined[text]["score"] + bm25_norm) / 2
        else:
            doc = Document(page_content=text, metadata=chunk_metas[idx])
            combined[text] = {"doc": doc, "score": bm25_norm * 0.5}
    ranked = sorted(combined.values(), key=lambda x: x["score"], reverse=True)
    return [(item["doc"], item["score"]) for item in ranked[:k]]

def rerank(query, results, top_n=3):
    scored = []
    for doc, score in results:
        text = doc.page_content.lower()
        query_words = query.lower().split()
        keyword_hits = sum(1 for w in query_words if w in text)
        combined_score = score * 0.7 + (keyword_hits / max(len(query_words), 1)) * 0.3
        scored.append((doc, combined_score))
    return sorted(scored, key=lambda x: x[1], reverse=True)[:top_n]

def generate_chat_docx(messages):
    doc = DocxDocument()
    title = doc.add_heading("AI Knowledge Assistant — Chat History", 0)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    doc.add_paragraph(f"Downloaded on: {time.strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")
    for msg in messages:
        if msg["role"] == "user":
            p = doc.add_paragraph()
            run = p.add_run(f"You: {msg['content']}")
            run.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        else:
            p = doc.add_paragraph()
            run = p.add_run(f"Assistant: {msg['content']}")
            run.font.size = Pt(10)
        doc.add_paragraph("")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

COMPANY_LIST = ["All Companies", "TCS", "Infosys", "Wipro", "Cognizant",
                "Google", "Tesla", "Amazon", "Microsoft", "Meta", "Johnson & Johnson"]

SUGGESTIONS = [
    "What is the leave policy in TCS?",
    "What is the maternity leave in Cognizant?",
    "What is the paternity leave in Wipro?",
    "What is the hybrid work policy in Wipro?",
    "What is the dress code policy?",
    "What is the exit process for employees?",
    "What is the vacation policy at Tesla?",
    "What is the parental leave at Google?",
    "What is the sick leave policy in TCS?",
    "What is the bereavement leave at J&J?",
    "What is the performance review system in Infosys?",
    "What is the learning platform in TCS?",
    "What is the Amazon return to office policy?",
    "What is the Microsoft parental leave policy?",
    "What is the J&J credo?",
    "What is the Career Choice programme at Amazon?",
    "What is Tesla return to office policy?",
    "How does Amazon RSU vesting work?",
    "What is the Microsoft growth mindset?",
    "What is the provident fund contribution in Cognizant?"
]

for key, val in {
    "messages": [], "feedback": {}, "dark_mode": False,
    "selected_question": None, "response_times": [],
    "total_questions": 0, "conversation_history": [],
    "agent_messages": [], "agent_history": [], "agent_selected": None,
    "agent_uploaded_db": None, "agent_uploaded_name": None
}.items():
    if key not in st.session_state:
        st.session_state[key] = val

if st.session_state.dark_mode:
    st.markdown("""
    <style>
    .stApp { background-color: #1a1a2e; color: #e0e0e0; }
    .stChatMessage { background-color: #16213e; }
    </style>
    """, unsafe_allow_html=True)

with st.sidebar:
    st.markdown("### About")
    st.write("AI-powered HR policy assistant for 10 companies.")
    st.markdown("---")
    dark = st.toggle("🌙 Dark Mode", value=st.session_state.dark_mode, key="dark_toggle")
    if dark != st.session_state.dark_mode:
        st.session_state.dark_mode = dark
        st.rerun()
    st.markdown("---")
    st.markdown("**Filter by Company:**")
    selected_company = st.selectbox(
        "Company", COMPANY_LIST,
        label_visibility="collapsed", key="company_filter"
    )
    st.markdown("---")
    st.markdown("**Response Language:**")
    lang_option = st.selectbox(
        "Language", ["English", "Tamil", "Hindi"],
        label_visibility="collapsed", key="lang_select"
    )
    lang_map = {"English": None, "Tamil": "ta", "Hindi": "hi"}
    target_lang = lang_map[lang_option]
    st.markdown("---")
    st.markdown("**Technical Features:**")
    use_hybrid = st.checkbox("Hybrid Search (BM25 + Vector)", value=True)
    use_rerank = st.checkbox("Reranker", value=True)
    use_memory = st.checkbox("Conversation Memory", value=True)
    use_expansion = st.checkbox("Query Expansion", value=False,
                                help="Slower but improves recall for vague questions")
    st.markdown("---")
    st.markdown("**Companies loaded:**")
    st.write("🇮🇳 TCS | Infosys | Wipro | Cognizant")
    st.write("🇺🇸 Google | Tesla | Amazon | Microsoft")
    st.write("🏥 Johnson & Johnson | Meta")
    st.markdown("---")
    st.metric("Documents loaded", count_docs())
    st.markdown("---")
    if st.session_state.total_questions > 0:
        st.markdown("**Chat Statistics:**")
        c1, c2 = st.columns(2)
        with c1:
            st.metric("Questions", st.session_state.total_questions)
        with c2:
            avg = round(sum(st.session_state.response_times) /
                        len(st.session_state.response_times), 1)
            st.metric("Avg time", f"{avg}s")
        st.markdown("---")
    if len(st.session_state.messages) > 0:
        docx_bytes = generate_chat_docx(st.session_state.messages)
        st.download_button(
            "📥 Download Chat as Word", data=docx_bytes,
            file_name="chat_history.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_docx"
        )
    if st.button("🗑️ Clear Chat History", key="clear_chat_btn"):
        for k in ["messages", "feedback", "response_times", "conversation_history"]:
            st.session_state[k] = [] if k != "feedback" else {}
        st.session_state.total_questions = 0
        st.rerun()

st.title("🤖 AI-Based Smart Knowledge Assistant")

tab1, tab2, tab3, tab4 = st.tabs(["💬 Chat", "⚖️ Compare Companies", "📊 Statistics", "🤖 AI Agent"])

# ════════════════════════════════
# TAB 1: CHAT
# ════════════════════════════════
with tab1:
    if len(st.session_state.messages) == 0:
        st.markdown("""
        <div class="welcome-box">
            <h3>Welcome! How can I help you today?</h3>
            <p>Ask questions about HR policies, leave, compensation, and more across 10 companies.</p>
        </div>
        """, unsafe_allow_html=True)
        st.markdown("**Try asking — click any question:**")
        example_questions = [
            "What is the leave policy in TCS?",
            "What is the maternity leave at Google?",
            "How does Amazon RSU vesting work?",
            "What is the dress code policy?",
            "What is Tesla's return to office policy?",
            "What is the J&J credo?"
        ]
        col1, col2 = st.columns(2)
        for i, q in enumerate(example_questions):
            with col1 if i % 2 == 0 else col2:
                if st.button(q, key=f"example_{i}", use_container_width=True):
                    st.session_state.selected_question = q
                    st.rerun()

    st.markdown("**Question suggestions:**")
    typed = st.text_input(
        "Suggestions", key="autocomplete_input",
        label_visibility="collapsed",
        placeholder="Start typing to filter suggestions..."
    )
    if typed and len(typed) >= 3:
        matches = [s for s in SUGGESTIONS if typed.lower() in s.lower()][:4]
        for s in matches:
            if st.button(s, key=f"sug_{s}", use_container_width=True):
                st.session_state.selected_question = s
                st.rerun()

    for i, msg in enumerate(st.session_state.messages):
        with st.chat_message(msg["role"]):
            st.write(msg["content"])
            if msg["role"] == "assistant":
                meta_html = f'Confidence: {msg.get("confidence","")}'
                if "response_time" in msg:
                    meta_html += f'<span class="time-badge">⏱ {msg["response_time"]}s</span>'
                if "techniques" in msg:
                    meta_html += f'<span class="tech-badge">🔬 {msg["techniques"]}</span>'
                st.markdown(meta_html, unsafe_allow_html=True)
                if "sources" in msg and msg["sources"]:
                    with st.expander("📄 Sources", expanded=False):
                        st.markdown(msg["sources"], unsafe_allow_html=True)
                fb_key = f"fb_{i}"
                if fb_key not in st.session_state.feedback:
                    c1, c2, c3 = st.columns([1, 1, 8])
                    with c1:
                        if st.button("👍", key=f"up_{i}"):
                            st.session_state.feedback[fb_key] = "up"
                            st.rerun()
                    with c2:
                        if st.button("👎", key=f"down_{i}"):
                            st.session_state.feedback[fb_key] = "down"
                            st.rerun()
                else:
                    if st.session_state.feedback[fb_key] == "up":
                        st.success("Thanks for the positive feedback! ✅")
                    else:
                        st.warning("Thanks! We will work on improving this answer.")

    if st.session_state.selected_question:
        question = st.session_state.selected_question
        st.session_state.selected_question = None
    else:
        question = st.chat_input("Ask a question about company policies...")

    if question:
        filtered_q = (f"{question} at {selected_company}"
                      if selected_company != "All Companies" else question)
        with st.chat_message("user"):
            st.write(question)
        st.session_state.messages.append({"role": "user", "content": question})

        with st.chat_message("assistant"):
            with st.spinner("Searching and generating answer..."):
                start_time = time.time()
                techniques = []
                queries = [filtered_q]
                if use_expansion:
                    queries = expand_query(filtered_q)
                    techniques.append("Query Expansion")
                all_results = {}
                for q in queries:
                    if use_hybrid:
                        results = hybrid_search(q, k=6)
                        if "Hybrid Search" not in techniques:
                            techniques.append("Hybrid Search")
                    else:
                        results = vector_db.similarity_search_with_score(q, k=6)
                    for doc, score in results:
                        key = doc.page_content
                        if key not in all_results or score > all_results[key][1]:
                            all_results[key] = (doc, score)
                results = list(all_results.values())
                if selected_company != "All Companies":
                    filtered = [
                        (doc, score) for doc, score in results
                        if selected_company.lower().replace(" & ", "").replace(" ", "")
                        in doc.metadata.get("source", "").lower().replace("_", "").replace("&", "")
                    ]
                    if filtered:
                        results = filtered
                if use_rerank:
                    results = rerank(question, results, top_n=4)
                    techniques.append("Reranker")
                else:
                    results = results[:4]
                context = "\n\n".join([doc.page_content for doc, _ in results])
                top_score = results[0][1] if results else 0
                if top_score > 0.7:
                    confidence = '<span class="confidence-high">● High confidence</span>'
                elif top_score > 0.4:
                    confidence = '<span class="confidence-med">● Medium confidence</span>'
                else:
                    confidence = '<span class="confidence-low">● Low confidence</span>'
                seen = []
                sources_html = ""
                for doc, score in results:
                    source = doc.metadata.get("source", "Unknown")
                    page = doc.metadata.get("page", "?")
                    company = get_company(source)
                    entry = f"{source}_{page}"
                    if entry not in seen:
                        seen.append(entry)
                        preview = doc.page_content[:180].replace("\n", " ").strip()
                        sources_html += (
                            f'<span class="company-tag">{company}</span>'
                            f' Page {page}<br>'
                            f'<div class="chunk-preview">"{preview}..."</div><br>'
                        )
                history = st.session_state.conversation_history if use_memory else []
                answer = get_llm_answer(question, context, history)
                if target_lang:
                    answer = translate_text(answer, target_lang)
                elapsed = round(time.time() - start_time, 1)
                st.session_state.response_times.append(elapsed)
                st.session_state.total_questions += 1
                if use_memory:
                    st.session_state.conversation_history.append({
                        "question": question, "answer": answer
                    })
                    if len(st.session_state.conversation_history) > 5:
                        st.session_state.conversation_history.pop(0)
                tech_str = " + ".join(list(dict.fromkeys(techniques)))

            st.write(answer)
            meta = f'Confidence: {confidence}<span class="time-badge">⏱ {elapsed}s</span>'
            if tech_str:
                meta += f'<span class="tech-badge">🔬 {tech_str}</span>'
            st.markdown(meta, unsafe_allow_html=True)
            with st.expander("📄 Sources", expanded=True):
                st.markdown(sources_html, unsafe_allow_html=True)

        st.session_state.messages.append({
            "role": "assistant", "content": answer,
            "confidence": confidence, "response_time": elapsed,
            "sources": sources_html, "techniques": tech_str
        })

# ════════════════════════════════
# TAB 2: COMPARE
# ════════════════════════════════
with tab2:
    st.markdown("### Compare HR Policies Side by Side")
    col1, col2 = st.columns(2)
    with col1:
        company_a = st.selectbox("Company A", COMPANY_LIST[1:], key="cmp_a")
    with col2:
        company_b = st.selectbox("Company B", COMPANY_LIST[1:], index=1, key="cmp_b")
    compare_topic = st.selectbox("Topic to compare", [
        "Leave policy", "Parental leave", "Work from home policy",
        "Performance review", "Compensation", "Bereavement leave",
        "Sick leave", "Learning and development"
    ], key="cmp_topic")
    if st.button("⚖️ Compare Now", key="compare_btn", use_container_width=True):
        if company_a == company_b:
            st.warning("Please select two different companies!")
        else:
            with st.spinner(f"Comparing {company_a} vs {company_b}..."):
                def get_company_answer(company, topic):
                    q = f"What is the {topic} at {company}?"
                    results = hybrid_search(q, k=4)
                    results = rerank(q, results, top_n=3)
                    context = "\n\n".join([doc.page_content for doc, _ in results])
                    return get_llm_answer(q, context)
                ans_a = get_company_answer(company_a, compare_topic)
                ans_b = get_company_answer(company_b, compare_topic)
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f"#### 🏢 {company_a}")
                st.info(ans_a)
            with col2:
                st.markdown(f"#### 🏢 {company_b}")
                st.info(ans_b)

# ════════════════════════════════
# TAB 3: STATISTICS
# ════════════════════════════════
with tab3:
    st.markdown("### Chat Statistics")
    if st.session_state.total_questions == 0:
        st.info("No questions asked yet. Start chatting to see statistics!")
    else:
        times = st.session_state.response_times
        thumbs_up = sum(1 for v in st.session_state.feedback.values() if v == "up")
        thumbs_down = sum(1 for v in st.session_state.feedback.values() if v == "down")
        c1, c2, c3 = st.columns(3)
        with c1:
            st.metric("Total Questions", st.session_state.total_questions)
        with c2:
            st.metric("Avg Response Time", f"{round(sum(times)/len(times),1)}s")
        with c3:
            st.metric("Fastest Response", f"{round(min(times),1)}s")
        c4, c5, c6 = st.columns(3)
        with c4:
            st.metric("Slowest Response", f"{round(max(times),1)}s")
        with c5:
            st.metric("👍 Positive", thumbs_up)
        with c6:
            st.metric("👎 Negative", thumbs_down)
        st.markdown("---")
        st.markdown("**Response Time per Question:**")
        st.line_chart(times)
        if thumbs_up + thumbs_down > 0:
            satisfaction = round((thumbs_up / (thumbs_up + thumbs_down)) * 100)
            st.markdown(f"**User Satisfaction: {satisfaction}%**")
            st.progress(satisfaction / 100)

# ════════════════════════════════
# TAB 4: AI AGENT
# ════════════════════════════════
with tab4:
    st.markdown("### 🤖 Agentic AI — Smart Multi-Step Reasoning")
    st.markdown("""
    The AI Agent thinks step by step and uses multiple tools automatically.
    You can also upload a new company PDF and the agent will search it too!
    """)
    st.markdown("---")

    st.markdown("#### 📄 Upload a Company PDF (Optional)")
    st.markdown("Upload any company HR policy PDF — the agent will include it in its search.")

    agent_uploaded_file = st.file_uploader(
        "Upload PDF for agent to use",
        type=["pdf"],
        key="agent_pdf_uploader"
    )

    if agent_uploaded_file is not None:
        if agent_uploaded_file.name != st.session_state.agent_uploaded_name:
            with st.spinner(f"Processing {agent_uploaded_file.name} for agent..."):
                try:
                    from langchain_text_splitters import RecursiveCharacterTextSplitter
                    from langchain_community.embeddings import HuggingFaceEmbeddings
                    from langchain_community.vectorstores import FAISS as FAISSUpload
                    from ocr_processor import process_uploaded_pdf

                    pdf_bytes = agent_uploaded_file.read()

                    st.info("🔍 Analysing document type...")
                    result, extracted_images = process_uploaded_pdf(
                        pdf_bytes, agent_uploaded_file.name
                    )
                    text = result["text"]

                    col_a, col_b, col_c = st.columns(3)
                    with col_a:
                        pdf_type_label = {
                            "text": "📝 Text PDF",
                            "scanned": "📸 Scanned PDF",
                            "mixed": "📄 Mixed PDF",
                            "unknown": "❓ Unknown"
                        }.get(result["pdf_type"], "📄 PDF")
                        st.metric("Document Type", pdf_type_label)
                    with col_b:
                        ocr_label = "✅ Yes" if result["ocr_used"] else "❌ No"
                        st.metric("OCR Used", ocr_label)
                    with col_c:
                        st.metric("Quality Score", f"{result['quality_score']}%")

                    if result["ocr_used"]:
                        st.warning(f"🔬 OpenCV preprocessing — {result['message']}")
                    else:
                        st.success(f"✅ {result['message']}")

                    if result["images_found"] > 0:
                        st.info(f"🖼️ Found {result['images_found']} chart/image region(s)")

                    if not text.strip():
                        st.error("Could not extract text from this PDF even with OCR.")
                    else:
                        splitter = RecursiveCharacterTextSplitter(
                            chunk_size=1000, chunk_overlap=200
                        )
                        chunks = splitter.create_documents(
                            [text],
                            metadatas=[{"source": agent_uploaded_file.name, "page": 0}]
                        )
                        upload_embeddings = HuggingFaceEmbeddings(
                            model_name="sentence-transformers/all-MiniLM-L6-v2"
                        )
                        agent_db = FAISSUpload.from_documents(chunks, upload_embeddings)
                        st.session_state.agent_uploaded_db = agent_db
                        st.session_state.agent_uploaded_name = agent_uploaded_file.name
                        from agent import set_uploaded_db
                        set_uploaded_db(agent_db, agent_uploaded_file.name)
                        st.success(f"✅ {agent_uploaded_file.name} loaded! Agent can now search this document.")

                except Exception as e:
                    st.error(f"Error: {str(e)}")

    if st.session_state.agent_uploaded_name:
        col1, col2 = st.columns([3, 1])
        with col1:
            st.info(f"📄 Agent has access to: **{st.session_state.agent_uploaded_name}** + all 10 company docs")
        with col2:
            if st.button("Remove PDF", key="remove_agent_pdf"):
                st.session_state.agent_uploaded_db = None
                st.session_state.agent_uploaded_name = None
                from agent import set_uploaded_db
                set_uploaded_db(None, None)
                st.rerun()
        search_mode = st.radio(
            "Agent should search:",
            ["All company docs + uploaded PDF",
             "Only uploaded PDF",
             "Only existing 10 companies"],
            horizontal=True,
            key="agent_search_mode"
        )
    else:
        search_mode = "Only existing 10 companies"

    st.markdown("---")
    st.markdown("**Try these complex questions:**")
    agent_examples = [
        "Compare leave policies of TCS and Google and tell me which is better",
        "Generate a full HR policy report for Amazon",
        "Which company has the best parental leave policy and why?",
        "How many total leave days does a TCS employee get per year?",
        "What should I consider before joining Tesla vs Microsoft?"
    ]
    if st.session_state.agent_uploaded_name:
        agent_examples.insert(0, "What is the leave policy in the uploaded document?")
        agent_examples.insert(1, "Compare the uploaded document with TCS policies")

    col1, col2 = st.columns(2)
    for i, q in enumerate(agent_examples[:6]):
        with col1 if i % 2 == 0 else col2:
            if st.button(q, key=f"agent_ex_{i}", use_container_width=True):
                st.session_state.agent_selected = q
                st.rerun()

    st.markdown("---")

    for msg in st.session_state.agent_messages:
        with st.chat_message(msg["role"]):
            st.write(msg["content"])
            if msg["role"] == "assistant" and "tools_used" in msg:
                if msg["tools_used"]:
                    tools_str = " → ".join(msg["tools_used"])
                    st.markdown(
                        f'<span class="tech-badge">🔧 Tools: {tools_str}</span>'
                        f'<span class="time-badge">⏱ {msg.get("response_time","?")}s</span>'
                        f'<span class="tech-badge">📋 {msg.get("steps",0)} steps</span>',
                        unsafe_allow_html=True
                    )

    if st.session_state.agent_selected:
        agent_question = st.session_state.agent_selected
        st.session_state.agent_selected = None
    else:
        agent_question = st.chat_input(
            "Ask the agent a complex question...",
            key="agent_input"
        )

    if agent_question:
        from agent import run_agent, set_uploaded_db

        if st.session_state.agent_uploaded_db:
            set_uploaded_db(
                st.session_state.agent_uploaded_db,
                st.session_state.agent_uploaded_name
            )

        use_uploaded_only = (search_mode == "Only uploaded PDF")

        with st.chat_message("user"):
            st.write(agent_question)
        st.session_state.agent_messages.append({
            "role": "user", "content": agent_question
        })

        with st.chat_message("assistant"):
            with st.spinner("Agent is thinking and using tools..."):
                start = time.time()
                result = run_agent(
                    agent_question,
                    st.session_state.agent_history,
                    use_uploaded=use_uploaded_only
                )
                elapsed = round(time.time() - start, 1)

            st.write(result["answer"])
            if result["tools_used"]:
                tools_str = " → ".join(result["tools_used"])
                st.markdown(
                    f'<span class="tech-badge">🔧 Tools: {tools_str}</span>'
                    f'<span class="time-badge">⏱ {elapsed}s</span>'
                    f'<span class="tech-badge">📋 {result["steps"]} steps</span>',
                    unsafe_allow_html=True
                )

        st.session_state.agent_messages.append({
            "role": "assistant",
            "content": result["answer"],
            "tools_used": result["tools_used"],
            "steps": result["steps"],
            "response_time": elapsed
        })

        st.session_state.agent_history.extend([
            {"role": "human", "content": agent_question},
            {"role": "assistant", "content": result["answer"]}
        ])
        if len(st.session_state.agent_history) > 10:
            st.session_state.agent_history = st.session_state.agent_history[-10:]
